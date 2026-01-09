# python-etl/pipelines/docx_pipeline.py

import io
import logging
import hashlib
from typing import List, Optional

from minio import Minio
import pandas as pd
from docx import Document   # python-docx

logger = logging.getLogger(__name__)


# ----------------- helpers: hashing & de-dup ----------------- #

def calculate_file_hash(data: bytes) -> str:
    """Stable SHA256 hash for duplicate detection."""
    return hashlib.sha256(data).hexdigest()


def is_duplicate(pg_conn, content_hash: str) -> bool:
    """
    Check in minio_data_catalog if this hash already exists.
    If yes, we can skip heavy processing.
    """
    if pg_conn is None:
        return False

    cursor = pg_conn.cursor()
    try:
        cursor.execute(
            """
            SELECT 1 FROM minio_data_catalog
            WHERE content_hash = %s
            LIMIT 1
            """,
            (content_hash,),
        )
        return cursor.fetchone() is not None
    finally:
        cursor.close()


# ----------------- helpers: text & table extraction ----------------- #

def _extract_text_from_docx_bytes(data: bytes) -> str:
    """
    Extract plain paragraph text from DOCX bytes using python-docx.
    """
    doc = Document(io.BytesIO(data))
    lines = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            lines.append(p.text)
    return "\n".join(lines)


def _extract_tables_from_docx_bytes(data: bytes) -> List[pd.DataFrame]:
    """
    Extract all tables from DOCX as list of pandas DataFrames.
    First row is treated as header if non-empty, else generic column names.
    """
    doc = Document(io.BytesIO(data))
    dfs: List[pd.DataFrame] = []

    for tbl_idx, tbl in enumerate(doc.tables):
        rows = []
        max_cols = max((len(r.cells) for r in tbl.rows), default=0)

        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < max_cols:
                cells += [""] * (max_cols - len(cells))
            rows.append(cells)

        if not rows:
            continue

        header = rows[0]
        data_rows = rows[1:] if len(rows) > 1 else []

        # Build column names (ensure not empty & not duplicated)
        cols = []
        used = {}
        for col in header:
            name = col.strip() or "col"
            if name in used:
                used[name] += 1
                name = f"{name}_{used[name]}"
            else:
                used[name] = 1
            cols.append(name)

        if not data_rows:  # table with only header
            df = pd.DataFrame(columns=cols)
        else:
            df = pd.DataFrame(data_rows, columns=cols)
            # Clean the dataframe
            df = df.dropna(how='all')  # Remove empty rows
            df = df.replace('', None)  # Replace empty strings with None

        if not df.empty:
            dfs.append(df)

    return dfs


def _ensure_unstructured_table(pg_conn):
    """
    Make sure unstructured_documents table exists with all required columns.
    """
    cursor = pg_conn.cursor()
    try:
        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS unstructured_documents (
                id SERIAL PRIMARY KEY,
                object_name  TEXT NOT NULL,
                file_type    TEXT NOT NULL,
                text_content TEXT,
                content_hash TEXT UNIQUE,
                created_at   TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        pg_conn.commit()
    except Exception:
        pg_conn.rollback()
        logger.exception("[docx] Failed ensuring unstructured_documents table")
        raise
    finally:
        cursor.close()


def _save_unstructured_doc(
    pg_conn,
    object_name: str,
    file_type: str,
    text_content: str,
    content_hash: str,
):
    """
    Save raw text into "unstructured_documents" for search/QA later.
    """
    if pg_conn is None:
        return

    _ensure_unstructured_table(pg_conn)
    
    cursor = pg_conn.cursor()
    try:
        cursor.execute(
            """
            INSERT INTO unstructured_documents
                (object_name, file_type, text_content, content_hash)
            VALUES (%s, %s, %s, %s)
            ON CONFLICT (content_hash) DO UPDATE
            SET text_content = EXCLUDED.text_content,
                object_name  = EXCLUDED.object_name
            """,
            (object_name, file_type, text_content, content_hash),
        )
        pg_conn.commit()
        logger.info(f"[docx] Saved text to unstructured_documents for {object_name}")
    except Exception:
        pg_conn.rollback()
        logger.exception("[docx] Failed inserting into unstructured_documents")
    finally:
        cursor.close()


def _process_extracted_table(
    minio_client,
    bucket_name: str,
    table_df: pd.DataFrame,
    table_key: str,
    pg_conn,
    catalog_updater,
):
    """
    Process a single extracted table from DOCX:
    1. Upload CSV to MinIO
    2. Create PostgreSQL table
    3. Update catalog
    """
    try:
        # 1. Upload CSV to MinIO
        csv_bytes = table_df.to_csv(index=False).encode("utf-8")
        minio_client.put_object(
            bucket_name,
            table_key,
            io.BytesIO(csv_bytes),
            length=len(csv_bytes),
            content_type="text/csv",
        )
        logger.info(f"[docx] Uploaded table CSV -> {table_key}")

        # 2. Create sanitized table name for PostgreSQL
        from pipelines.structured_pipeline import sanitize_table_name, sanitize_column_name, infer_postgres_type
        
        table_name = sanitize_table_name(table_key)
        
        # Sanitize column names
        table_df.columns = [sanitize_column_name(col) for col in table_df.columns]
        
        # Handle duplicate column names
        cols = pd.Series(table_df.columns)
        for dup in cols[cols.duplicated()].unique():
            cols[cols == dup] = [f"{dup}_{i}" if i != 0 else dup for i in range(sum(cols == dup))]
        table_df.columns = cols
        
        # 3. Load to PostgreSQL
        cursor = pg_conn.cursor()
        
        # Drop existing table
        cursor.execute(f'DROP TABLE IF EXISTS "{table_name}"')
        
        # Create table with inferred types
        column_defs = []
        for col in table_df.columns:
            pg_type = infer_postgres_type(table_df[col])
            column_defs.append(f'"{col}" {pg_type}')
        
        create_sql = f'CREATE TABLE "{table_name}" ({", ".join(column_defs)})'
        cursor.execute(create_sql)
        logger.info(f"[docx] Created table: {table_name}")
        
        # Bulk insert
        buffer = io.StringIO()
        table_df.to_csv(buffer, index=False, header=False, sep='\t', na_rep='\\N')
        buffer.seek(0)
        
        cursor.copy_expert(
            f'COPY "{table_name}" FROM STDIN WITH (FORMAT CSV, DELIMITER E\'\\t\', NULL \'\\N\')',
            buffer
        )
        
        pg_conn.commit()
        logger.info(f"[docx] Loaded {len(table_df)} rows into {table_name}")
        cursor.close()

        # 4. Update catalog for the extracted table CSV
        catalog_updater(
            object_name=table_key,
            object_size=len(csv_bytes),
            file_format='csv',
            row_count=len(table_df),
            text_extracted=False,
            content_hash=None
        )

    except Exception as e:
        logger.exception(f"[docx] Failed processing table {table_key}: {e}")
        if pg_conn:
            pg_conn.rollback()


# ----------------- main entrypoint for Airflow/Lakehouse ----------------- #

def process_minio_object(
    minio_client: Minio,
    bucket_name: str,
    object_name: str,
    pg_conn,
    catalog_updater,
):
    """
    Main function used by LakehouseETL.run_pipeline_for_object.

    Steps:
    1. Download DOC/DOCX from MinIO.
    2. Compute SHA256 hash & check duplicate.
    3. Extract text (DOCX fully, DOC best-effort).
    4. Save text to:
        - MinIO: processed/unstructured/text-extracted/<file>.txt
        - Postgres: unstructured_documents table
    5. Extract tables from DOCX:
        - Upload each as CSV to processed/structured/docx-tables/<file>_table_X.csv
        - Create PostgreSQL table for each
        - Update catalog for each table
    6. Update minio_data_catalog for original DOCX file via catalog_updater(...).
    """

    logger.info(f"[docx] Processing {object_name}")

    # ---- 1. download bytes from MinIO ---- #
    response = minio_client.get_object(bucket_name, object_name)
    data = response.read()
    response.close()
    response.release_conn()

    file_size = len(data)
    file_hash = calculate_file_hash(data)
    ext = object_name.rsplit(".", 1)[-1].lower()
    root_name = object_name.split("/")[-1].rsplit(".", 1)[0]
    file_type = ext  # 'doc' or 'docx'

    # ---- 2. duplicate check ---- #
    if is_duplicate(pg_conn, file_hash):
        logger.info(
            f"[docx] Duplicate detected via content_hash, "
            f"skipping heavy processing: {object_name}"
        )
        # still make sure catalog has an entry for this object
        try:
            catalog_updater(
                object_name=object_name,
                object_size=file_size,
                file_format=file_type,
                row_count=0,
                text_extracted=False,
                content_hash=file_hash,
            )
        except Exception:
            logger.exception("[docx] Failed catalog update for duplicate")
        return

    # ---- 3. text extraction ---- #
    text: str = ""
    tables: List[pd.DataFrame] = []
    
    try:
        if ext == "docx":
            logger.info(f"[docx] Extracting text from DOCX...")
            text = _extract_text_from_docx_bytes(data)
            logger.info(f"[docx] Extracted {len(text)} characters of text")
            
            logger.info(f"[docx] Extracting tables from DOCX...")
            tables = _extract_tables_from_docx_bytes(data)
            logger.info(f"[docx] Found {len(tables)} tables")
            
        elif ext == "doc":
            # OLD binary format – python-docx cannot handle it.
            # For now: we mark it as unsupported for deep parsing but
            # avoid crashing the pipeline.
            logger.warning(
                "[docx] .doc format is not fully supported by python-docx. "
                "Text/tables will not be extracted; only catalog entry is updated. "
                "Consider converting to .docx format for full processing."
            )
        else:
            logger.warning(f"[docx] Unexpected extension {ext}, treating as binary blob.")
            
    except Exception as e:
        logger.exception(f"[docx] Error extracting content from {object_name}: {e}")

    # ---- 4. save raw text to MinIO + Postgres ---- #
    if text and text.strip():
        text_bytes = text.encode("utf-8")
        text_key = f"processed/unstructured/text-extracted/{root_name}.txt"

        minio_client.put_object(
            bucket_name,
            text_key,
            io.BytesIO(text_bytes),
            length=len(text_bytes),
            content_type="text/plain",
        )
        logger.info(f"[docx] Uploaded extracted text -> {text_key}")

        _save_unstructured_doc(
            pg_conn=pg_conn,
            object_name=object_name,
            file_type=file_type,
            text_content=text,
            content_hash=file_hash,
        )
    else:
        logger.warning(f"[docx] No text extracted from {object_name}")

    # ---- 5. extract + process tables (DOCX only) ---- #
    total_rows = 0
    
    if tables:
        logger.info(f"[docx] Processing {len(tables)} tables...")
        
        for idx, df in enumerate(tables, start=1):
            if df.empty:
                logger.debug(f"[docx] Skipping empty table {idx}")
                continue
                
            table_key = (
                f"processed/structured/docx-tables/{root_name}_table_{idx}.csv"
            )

            # Process this table
            _process_extracted_table(
                minio_client=minio_client,
                bucket_name=bucket_name,
                table_df=df,
                table_key=table_key,
                pg_conn=pg_conn,
                catalog_updater=catalog_updater,
            )

            total_rows += len(df)
            logger.info(
                f"[docx] Processed table {idx}/{len(tables)}: "
                f"{len(df)} rows, {len(df.columns)} columns"
            )

    # ---- 6. catalog update for original DOC/DOCX file ---- #
    try:
        # NEW: Metadata (after Document)
        doc = Document(io.BytesIO(data))
        props = doc.core_properties
        doc_metadata = {
            'author': props.author,
            'created': props.created.isoformat() if props.created else None,
            'modified': props.modified.isoformat() if props.modified else None,
            'title': props.title,
            'subject': props.subject,
            'category': props.category,
            'comments': props.comments,
            'table_count': len(tables),
            'paragraph_count': len(doc.paragraphs)
        }

        catalog_updater(
            object_name=object_name,
            object_size=file_size,
            file_format=file_type,
            row_count=total_rows,
            text_extracted=bool(text and text.strip()),
            content_hash=file_hash,
            metadata=doc_metadata  # NEW
        )
        logger.info(f"[docx] Updated catalog for {object_name}")
    except Exception as e:
        logger.exception(f"[docx] Failed catalog update for original document: {e}")
        # don't re-raise; pipeline has already done the heavy work

    logger.info(
        f"[docx] ✅ Completed processing {object_name} | "
        f"tables={len(tables)} | total_rows={total_rows} | text_chars={len(text)}"
    )